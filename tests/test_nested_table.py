"""Test nested table formatting (indent + line spacing)."""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from backend.formatter.engine import format_docx
from shared.schemas import ProfileConfig


def collect_all_cells(table, cells=None):
    """Recursively collect all cells including nested tables."""
    from docx.table import Table
    if cells is None:
        cells = []
    for row in table.rows:
        for cell in row.cells:
            cells.append(cell)
            for tbl_elem in cell._tc.findall(qn("w:tbl")):
                nested = Table(tbl_elem, None)
                collect_all_cells(nested, cells)
    return cells


def main():
    test_input = Path(__file__).parent / "_test_nested_input.docx"
    test_output = Path(__file__).parent / "_test_nested_output.docx"

    # Create doc: outer table with a nested table inside cell(0,0)
    doc = Document()
    style = doc.styles['Normal']
    style.paragraph_format.first_line_indent = Pt(21)  # 2-char-ish indent to inherit

    outer = doc.add_table(rows=2, cols=2)
    outer.cell(0, 1).text = "outer cell text"
    outer.cell(1, 0).text = "outer cell 2"

    # Add nested table inside outer cell(0,0)
    inner_cell = outer.cell(0, 0)
    inner_tbl = inner_cell.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            inner_tbl.cell(r, c).text = f"nested {r}-{c}"

    doc.save(str(test_input))
    print(f"Input created with nested table")

    # Format with table indent=none, line_spacing=2.0
    profile = ProfileConfig()
    profile.table.indent_type = "none"
    profile.table.line_spacing = 2.0
    profile.table.line_spacing_mode = "multiple"

    success, msg, _ = format_docx(str(test_input), profile, output_path=str(test_output))
    print(f"Format: success={success}")

    # Verify ALL cells (outer + nested)
    doc_out = Document(str(test_output))
    total = 0
    errors = []
    for table in doc_out.tables:
        for cell in collect_all_cells(table):
            for para in cell.paragraphs:
                total += 1
                ppr = para._element.find(qn("w:pPr"))
                if ppr is None:
                    errors.append(f"para '{para.text[:15]}': NO pPr")
                    continue
                ind = ppr.find(qn("w:ind"))
                if ind is None:
                    errors.append(f"para '{para.text[:15]}': NO w:ind")
                else:
                    for attr, val in ind.attrib.items():
                        if val != "0":
                            tag = attr.split("}")[-1]
                            errors.append(f"para '{para.text[:15]}': {tag}={val}")
                spacing = ppr.find(qn("w:spacing"))
                if spacing is None:
                    errors.append(f"para '{para.text[:15]}': NO w:spacing")
                else:
                    line = spacing.get(qn("w:line"))
                    if line != "480":
                        errors.append(f"para '{para.text[:15]}': line={line} (expect 480)")

    print(f"\nChecked {total} cell paragraphs (outer + nested)")
    if errors:
        print(f"{len(errors)} ERRORS:")
        for e in errors:
            print(f"  {e}")
        sys.exit(1)
    else:
        print("ALL PASS: indent=0 and line=480 (2x) in every cell incl. nested")


if __name__ == "__main__":
    main()

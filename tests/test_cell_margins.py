"""Test cell margin (H/V) application including override of existing tcMar."""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from pathlib import Path
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from backend.formatter.engine import format_docx
from shared.schemas import ProfileConfig

HERE = Path(__file__).parent


def main():
    src = HERE / "_margin_input.docx"
    out = HERE / "_margin_output.docx"

    # Create doc: table whose cells carry their OWN tcMar (simulates real docs)
    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    for row in tbl.rows:
        for cell in row.cells:
            cell.text = "内容"
            tc = cell._tc
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                tc.insert(0, tc_pr)
            tc_pr.append(parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'<w:left w:w="600" w:type="dxa"/>'
                f'<w:right w:w="600" w:type="dxa"/>'
                f'</w:tcMar>'))
    doc.save(str(src))

    # Format: H margin 0.29cm, V margin 0.37cm
    profile = ProfileConfig()
    profile.table.cell_margin_h = 0.29
    profile.table.cell_margin_h_unit = "cm"
    profile.table.cell_margin_v = 0.37
    profile.table.cell_margin_v_unit = "cm"
    # Header bg color to also test shd/tcMar ordering
    profile.table.header_bg_color = "#D9E2F3"

    ok_fmt, msg, _ = format_docx(str(src), profile, output_path=str(out))
    print("format:", ok_fmt)

    exp_h = str(int(0.29 * 567))  # 164
    exp_v = str(int(0.37 * 567))  # 209
    print(f"expect H={exp_h} twips, V={exp_v} twips")

    doc2 = Document(str(out))
    errors = []

    # OOXML schema order for tcPr children (subset we write)
    tcpr_order = ["w:tcW", "w:gridSpan", "w:tcBorders", "w:shd", "w:noWrap",
                  "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark"]

    for t_i, table in enumerate(doc2.tables):
        # table-level tblCellMar position: must be before tblLook
        tbl_pr = table._tbl.tblPr
        children = [c.tag.split("}")[-1] for c in tbl_pr]
        if "tblCellMar" not in children:
            errors.append("tblPr missing tblCellMar")
        elif "tblLook" in children and children.index("tblCellMar") > children.index("tblLook"):
            errors.append(f"tblCellMar AFTER tblLook: {children}")

        for r_i, row in enumerate(table.rows):
            for c_i, cell in enumerate(row.cells):
                tc_pr = cell._tc.find(qn("w:tcPr"))
                if tc_pr is None:
                    errors.append(f"[{r_i},{c_i}] no tcPr")
                    continue
                tc_mar = tc_pr.find(qn("w:tcMar"))
                if tc_mar is None:
                    errors.append(f"[{r_i},{c_i}] no tcMar")
                    continue
                vals = {}
                for side in ("top", "left", "bottom", "right"):
                    el = tc_mar.find(qn(f"w:{side}"))
                    vals[side] = el.get(qn("w:w")) if el is not None else None
                if vals["left"] != exp_h or vals["right"] != exp_h:
                    errors.append(f"[{r_i},{c_i}] H margin wrong: {vals}")
                if vals["top"] != exp_v or vals["bottom"] != exp_v:
                    errors.append(f"[{r_i},{c_i}] V margin wrong: {vals}")

                # element order check
                tags = ["w:" + c.tag.split("}")[-1] for c in tc_pr]
                known = [t for t in tags if t in tcpr_order]
                sorted_known = sorted(known, key=tcpr_order.index)
                if known != sorted_known:
                    errors.append(f"[{r_i},{c_i}] tcPr order invalid: {tags}")

    if errors:
        print(f"{len(errors)} ERRORS:")
        for e in errors:
            print(" ", e)
        sys.exit(1)
    print("PASS: all cells tcMar H=0.29cm V=0.37cm, tblCellMar ordered, tcPr order valid")


if __name__ == "__main__":
    main()

"""Test with a document that has default Chinese formatting (2-char indent)."""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

from backend.formatter.engine import format_docx
from backend.formatter.table import _set_cell_alignment, _apply_cell_line_spacing
from backend.formatter.data_model import TableConfig as DmTableConfig
from shared.schemas import ProfileConfig


def test_direct_cell_functions():
    """Test _set_cell_alignment and _apply_cell_line_spacing directly on a cell."""
    print("=== Direct function test ===")

    doc = Document()

    # Set Normal style to have first-line indent and fixed line spacing
    style = doc.styles['Normal']
    pf = style.paragraph_format
    pf.first_line_indent = Pt(21)  # ~2 chars
    pf.line_spacing = Pt(18)  # fixed 18pt
    pf.line_spacing_rule = 4  # EXACTLY
    print(f"Normal style indent: {pf.first_line_indent}")
    print(f"Normal style line_spacing: {pf.line_spacing}")
    print(f"Normal style line_spacing_rule: {pf.line_spacing_rule}")

    # Add a paragraph to see if Normal style applies
    p = doc.add_paragraph("测试正文")
    print(f"Para indent: {p.paragraph_format.first_line_indent}")
    print(f"Para line_spacing: {p.paragraph_format.line_spacing}")

    # Add table
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "测试表格"

    # Inspect cell paragraph before formatting
    for para in cell.paragraphs:
        pf = para.paragraph_format
        ppr = para._element.find(qn("w:pPr"))
        has_ind = ppr is not None and ppr.find(qn("w:ind")) is not None
        has_sp = ppr is not None and ppr.find(qn("w:spacing")) is not None
        print(f"\nBefore: has paragraph-level ind: {has_ind}, has paragraph-level spacing: {has_sp}")
        print(f"  line_spacing_rule (inherited): {pf.line_spacing_rule}")

    # Apply table formatting with indent=none, line_spacing=2.0 multiple
    config = DmTableConfig()
    config.indent_type = "none"
    config.indent_value = 0.0
    config.indent_unit = "字符"
    config.line_spacing = 2.0
    config.line_spacing_mode = "multiple"

    _set_cell_alignment(cell, config)
    _apply_cell_line_spacing(cell, config)

    # Inspect after
    for para in cell.paragraphs:
        ppr = para._element.find(qn("w:pPr"))
        ind = ppr.find(qn("w:ind"))
        spacing = ppr.find(qn("w:spacing"))

        print(f"\nAfter _set_cell_alignment + _apply_cell_line_spacing:")
        if ind is not None:
            print(f"  w:ind attributes: {dict(ind.attrib)}")
            # Check if any non-zero
            for k, v in ind.attrib.items():
                tag = k.split("}")[-1] if "}" in k else k
                if v != "0":
                    print(f"  *** NON-ZERO: {tag}={v} ***")
        else:
            print(f"  NO w:ind element!")

        if spacing is not None:
            line = spacing.get(qn("w:line"))
            rule = spacing.get(qn("w:lineRule"))
            print(f"  w:spacing w:line={line} w:lineRule={rule}")
        else:
            print(f"  NO w:spacing element!")


def test_full_format_with_body_indent():
    """Test that body indent doesn't leak into table cells."""
    print("\n=== Full format test with body indent ===")

    test_input = Path(__file__).parent / "_test_body_indent_input.docx"

    # Create doc with a body paragraph and a table
    doc = Document()

    # Body paragraph with 2-char indent
    body_para = doc.add_paragraph("正文内容——应该有首行缩进")
    body_para.paragraph_format.first_line_indent = Pt(21)

    # Table with content
    table = doc.add_table(rows=2, cols=2)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"表格内容 {i+1}-{j+1}"

    doc.save(str(test_input))
    print(f"Input created: {test_input}")

    # Profile with body indent=first_line 2char, table indent=none
    profile = ProfileConfig()
    profile.body.indent_type = "firstLine"
    profile.body.indent_value = 2.0
    profile.body.indent_unit = "字符"
    profile.body.line_spacing = 1.5
    profile.body.line_spacing_mode = "multiple"
    profile.table.indent_type = "none"
    profile.table.indent_value = 0.0
    profile.table.line_spacing = 2.0
    profile.table.line_spacing_mode = "multiple"

    test_output = Path(__file__).parent / "_test_body_indent_output.docx"
    success, msg, out_path = format_docx(str(test_input), profile, output_path=str(test_output))
    print(f"Format: success={success}, msg={msg}")

    # Check output
    doc_out = Document(str(test_output))

    # Check body paragraph
    for para in doc_out.paragraphs:
        ppr = para._element.find(qn("w:pPr"))
        if ppr is not None:
            ind = ppr.find(qn("w:ind"))
            if ind is not None:
                fl = ind.get(qn("w:firstLine"))
                flc = ind.get(qn("w:firstLineChars"))
                print(f"Body para '{para.text[:20]}...': firstLine={fl}, firstLineChars={flc}")
            else:
                print(f"Body para '{para.text[:20]}...': NO indent element")

    # Check table cells
    errors = []
    for table in doc_out.tables:
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    ppr = para._element.find(qn("w:pPr"))
                    if ppr is None:
                        errors.append(f"[{row_idx},{col_idx}] para[{para_idx}]: NO pPr")
                        continue
                    ind = ppr.find(qn("w:ind"))
                    if ind is None:
                        errors.append(f"[{row_idx},{col_idx}] para[{para_idx}]: NO w:ind")
                        continue
                    # Check for non-zero indent
                    for attr, val in ind.attrib.items():
                        tag = attr.split("}")[-1] if "}" in attr else attr
                        if val != "0":
                            errors.append(f"[{row_idx},{col_idx}] para[{para_idx}]: {tag}={val} (non-zero!)")

                    spacing = ppr.find(qn("w:spacing"))
                    if spacing is not None:
                        line = spacing.get(qn("w:line"))
                        if line:
                            line_val = int(line)
                            if abs(line_val - 480) > 1:
                                errors.append(f"[{row_idx},{col_idx}] para[{para_idx}]: line={line_val} (not 2x!)")

    if errors:
        print(f"\n{len(errors)} ERRORS FOUND:")
        for e in errors:
            print(f"  {e}")
    else:
        print("\nAll table cells: indent=0, line_spacing=2x ✓")


if __name__ == "__main__":
    test_direct_cell_functions()
    test_full_format_with_body_indent()

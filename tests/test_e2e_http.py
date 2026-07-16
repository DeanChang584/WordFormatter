"""E2E test: simulate the frontend's HTTP format request against a live backend."""
import json
import sys
import time
import urllib.request
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

BASE = "http://127.0.0.1:8765/api"
HERE = Path(__file__).parent


def post(path, body):
    req = urllib.request.Request(
        BASE + path, data=json.dumps(body).encode(),
        headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req) as r:
        return json.loads(r.read())


def get(path):
    with urllib.request.urlopen(BASE + path) as r:
        return json.loads(r.read())


def main():
    # 1. Create test doc: table with 1.3x line spacing + 2-char indent
    src = HERE / "_e2e_input.docx"
    doc = Document()
    style = doc.styles["Normal"]
    style.paragraph_format.first_line_indent = Pt(21)
    tbl = doc.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            cell = tbl.cell(r, c)
            cell.text = f"cell {r}-{c}"
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.3  # original 1.3x
    doc.save(str(src))

    # 2. Send format request EXACTLY like the frontend (camelCase JSON)
    out_dir = HERE / "_e2e_out"
    out_dir.mkdir(exist_ok=True)
    profile = {
        "table": {
            "indentType": "none",
            "indentValue": 0.0,
            "indentUnit": "字符",
            "lineSpacing": 2.0,
            "lineSpacingMode": "multiple",
        }
    }
    resp = post("/format/start", {
        "files": [str(src)],
        "profile": profile,
        "outputDir": str(out_dir),
    })
    task_id = resp["data"]["taskId"]
    print(f"Task: {task_id}")

    # 3. Poll until done
    for _ in range(60):
        st = get(f"/format/status/{task_id}")["data"]
        if st.get("state") in ("done", "finished", "completed", "error", "cancelled"):
            break
        time.sleep(0.5)
    print(f"State: {st.get('state')}")

    # 4. Verify output
    outputs = list(out_dir.glob("*.docx"))
    if not outputs:
        print("FAIL: no output file")
        sys.exit(1)
    out_file = max(outputs, key=lambda p: p.stat().st_mtime)
    print(f"Output: {out_file.name}")

    doc2 = Document(str(out_file))
    ok = True
    for table in doc2.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    ppr = para._element.find(qn("w:pPr"))
                    ind = ppr.find(qn("w:ind")) if ppr is not None else None
                    spacing = ppr.find(qn("w:spacing")) if ppr is not None else None
                    ind_ok = ind is not None and all(v == "0" for v in ind.attrib.values())
                    line = spacing.get(qn("w:line")) if spacing is not None else None
                    sp_ok = line == "480"
                    if not ind_ok or not sp_ok:
                        ok = False
                        print(f"  BAD '{para.text[:10]}': ind_ok={ind_ok}, line={line}")
    print("E2E RESULT:", "PASS - indent cleared, line spacing = 2x (480)" if ok else "FAIL")
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()

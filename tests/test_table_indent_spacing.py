"""Test that table cell indent and line spacing are correctly applied."""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from docx.oxml.ns import qn

from backend.formatter.engine import format_docx
from shared.schemas import ProfileConfig

def create_test_doc(path: str):
    """Create a .docx with a table where cells have 2-char first-line indent."""
    doc = Document()

    # Set Normal style to have first-line indent of 2 chars
    style = doc.styles['Normal']
    style.paragraph_format.first_line_indent = Pt(21)  # ~2 chars at 10.5pt

    # Add a table
    table = doc.add_table(rows=2, cols=2)
    for row in table.rows:
        for cell in row.cells:
            cell.text = "测试文本内容"

    doc.save(path)
    print(f"Test doc created: {path}")

def inspect_table_xml(docx_path: str):
    """Inspect the XML of table cells for indent and line spacing."""
    doc = Document(docx_path)
    for i, table in enumerate(doc.tables):
        for j, row in enumerate(table.rows):
            for k, cell in enumerate(row.cells):
                for l, para in enumerate(cell.paragraphs):
                    ppr = para._element.find(qn("w:pPr"))
                    if ppr is not None:
                        ind = ppr.find(qn("w:ind"))
                        spacing = ppr.find(qn("w:spacing"))
                        ind_str = "none"
                        if ind is not None:
                            parts = []
                            for attr, val in sorted(ind.attrib.items()):
                                tag = attr.split("}")[-1] if "}" in attr else attr
                                if val != "0":
                                    parts.append(f"{tag}={val}")
                            ind_str = ", ".join(parts) if parts else "all-zero"

                        sp_str = "none"
                        if spacing is not None:
                            line = spacing.get(qn("w:line"))
                            rule = spacing.get(qn("w:lineRule"))
                            if line:
                                line_val = int(line)
                                if rule in ("exact", "atLeast"):
                                    sp_str = f"line={line_val}EMU ({line_val/12700:.1f}pt) rule={rule}"
                                else:
                                    sp_str = f"line={line_val} (×{line_val/240:.1f}) rule={rule}"

                        print(f"  Table[{i}] Row[{j}] Cell[{k}] Para[{l}]: indent=[{ind_str}] spacing=[{sp_str}]")

def main():
    test_input = Path(__file__).parent / "_test_table_input.docx"
    test_output = Path(__file__).parent / "_test_table_output.docx"

    # Create test doc
    create_test_doc(str(test_input))

    # Inspect original
    print("\n=== Original document table XML ===")
    inspect_table_xml(str(test_input))

    # Create a profile with table indent=none, table line_spacing=2.0
    profile = ProfileConfig()
    profile.table.indent_type = "none"
    profile.table.indent_value = 0.0
    profile.table.line_spacing = 2.0
    profile.table.line_spacing_mode = "multiple"

    # Format
    success, msg, out_path = format_docx(str(test_input), profile, output_path=str(test_output))
    print(f"\nFormat result: success={success}, msg={msg}")

    # Inspect output
    print("\n=== Formatted document table XML ===")
    inspect_table_xml(str(test_output))

    # Check that indent is all zeros
    doc = Document(str(test_output))
    all_zero = True
    all_spacing_2x = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    ppr = para._element.find(qn("w:pPr"))
                    if ppr is not None:
                        ind = ppr.find(qn("w:ind"))
                        if ind is not None:
                            for attr, val in ind.attrib.items():
                                if val != "0":
                                    tag = attr.split("}")[-1] if "}" in attr else attr
                                    print(f"  NON-ZERO INDENT: {tag}={val}")
                                    all_zero = False
                        else:
                            print(f"  NO w:ind element found!")
                            all_zero = False

                        spacing = ppr.find(qn("w:spacing"))
                        if spacing is not None:
                            line = spacing.get(qn("w:line"))
                            if line:
                                line_val = int(line)
                                rule = spacing.get(qn("w:lineRule"))
                                if rule in ("exact", "atLeast"):
                                    print(f"  LINE SPACING not multiple: line={line_val} rule={rule}")
                                    all_spacing_2x = False
                                elif abs(line_val - 480) > 1:
                                    print(f"  LINE SPACING not 2x: line={line_val} rule={rule}")
                                    all_spacing_2x = False
                        else:
                            print(f"  NO w:spacing element found!")
                            all_spacing_2x = False

    print(f"\n=== RESULTS ===")
    print(f"Indent all zero: {all_zero}")
    print(f"Line spacing 2x: {all_spacing_2x}")

if __name__ == "__main__":
    main()

"""
验证首行缩进修复：旧值清除和 w:firstLine 写入正确。
"""
from docx import Document
from docx.oxml.ns import qn
from backend.formatter.paragraph import (
    apply_first_line_indent, apply_paragraph_format,
)
from backend.formatter.data_model import BodyConfig


def test_first_line_indent_2_char():
    """2字符缩进 → w:firstLine=480 twips + w:firstLineChars=200

    同时写入两个属性：
      - w:firstLine（twips）控制实际渲染，兼容 WPS
      - w:firstLineChars（1/100 字符）供 Word UI 显示"2字符"
    """
    doc = Document()
    para = doc.add_paragraph("测试段落")
    apply_first_line_indent(para, value=2.0, unit="字符", font_size_pt=12.0)

    pPr = para._element.find(qn("w:pPr"))
    ind = pPr.find(qn("w:ind"))
    assert ind is not None
    first_line = ind.get(qn("w:firstLine"))
    assert first_line == "480", f"Expected 480, got {first_line}"
    first_line_chars = ind.get(qn("w:firstLineChars"))
    assert first_line_chars == "200", f"Expected 200, got {first_line_chars}"
    # 确保旧属性被清理
    assert qn("w:hanging") not in ind.attrib
    assert qn("w:hangingChars") not in ind.attrib


def test_first_line_indent_update():
    """重复调用应覆盖旧值，同时更新 firstLine 和 firstLineChars"""
    doc = Document()
    para = doc.add_paragraph("测试段落")
    # 先设置 2 字符
    apply_first_line_indent(para, value=2.0, unit="字符", font_size_pt=12.0)
    # 再改为 1 字符
    apply_first_line_indent(para, value=1.0, unit="字符", font_size_pt=12.0)

    pPr = para._element.find(qn("w:pPr"))
    ind = pPr.find(qn("w:ind"))
    first_line = ind.get(qn("w:firstLine"))
    assert first_line == "240", f"Expected 240, got {first_line}"
    first_line_chars = ind.get(qn("w:firstLineChars"))
    assert first_line_chars == "100", f"Expected 100, got {first_line_chars}"


def test_first_line_indent_legacy_removal():
    """模拟旧文档残留 firstLineChars，确保被清除并重写"""
    doc = Document()
    para = doc.add_paragraph("测试段落")
    pPr = para._element.get_or_add_pPr()
    # 手工构造包含 firstLineChars 和 firstLine 的 w:ind 元素
    ind_xml = (
        '<w:ind xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:firstLineChars="200" w:firstLine="300" />'
    )
    from docx.oxml import parse_xml
    ind = parse_xml(ind_xml)
    pPr.append(ind)

    apply_first_line_indent(para, value=2.0, unit="字符", font_size_pt=12.0)

    # firstLineChars 应被清除并写入新值，firstLine 应更新为新值
    ind = pPr.find(qn("w:ind"))
    assert ind.get(qn("w:firstLine")) == "480", f"Expected 480, got {ind.get(qn('w:firstLine'))}"
    assert ind.get(qn("w:firstLineChars")) == "200", f"Expected 200, got {ind.get(qn('w:firstLineChars'))}"


def test_first_line_indent_no_chars_when_cm():
    """厘米单位不应写入 w:firstLineChars"""
    doc = Document()
    para = doc.add_paragraph("测试段落")
    apply_first_line_indent(para, value=0.85, unit="cm")

    pPr = para._element.find(qn("w:pPr"))
    ind = pPr.find(qn("w:ind"))
    # 厘米单位时 firstLineChars 设为 0
    assert ind.get(qn("w:firstLineChars")) == "0", f"Expected 0, got {ind.get(qn('w:firstLineChars'))}"


def test_paragraph_format_full_flow():
    """通过 apply_paragraph_format 完整流程验证 indent_value=2.0"""
    config = BodyConfig(
        indent_type="firstLine",
        indent_value=2.0,
        indent_unit="字符",
        font_size=12.0,
    )
    doc = Document()
    para = doc.add_paragraph("测试段落")
    apply_paragraph_format(para, config)

    pPr = para._element.find(qn("w:pPr"))
    ind = pPr.find(qn("w:ind"))
    assert ind is not None
    first_line = ind.get(qn("w:firstLine"))
    assert first_line == "480", f"Expected 480, got {first_line}"
    first_line_chars = ind.get(qn("w:firstLineChars"))
    assert first_line_chars == "200", f"Expected 200, got {first_line_chars}"
"""
图片规则 — 自动压缩大图片（质量 + 尺寸）+ 图片尺寸设置

检测文档中所有嵌入图片（包括页眉/页脚中的图片），
对超过限制的图片进行缩放和压缩。

接口:
    compress_images(doc, config) -> ImageCompressResult
    apply_image_size(doc, config) -> None

依赖:
    Pillow>=10.0.0 （optional — 未安装则静默跳过）
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from docx import Document
    from docx.opc.part import Part

from .data_model import PictureConfig

logger = logging.getLogger("backend.formatter.image")


# ============================================================
# 结果封装
# ============================================================


@dataclass
class ImageCompressResult:
    """图片压缩处理结果统计。"""

    processed: int = 0          # 实际压缩的图片数
    before_bytes: int = 0       # 压缩前总字节
    after_bytes: int = 0        # 压缩后总字节
    skipped: int = 0            # 无需压缩的图片数
    errors: list[str] = field(default_factory=list)  # 处理失败记录

    @property
    def saved_bytes(self) -> int:
        """节省的字节数。"""
        return self.before_bytes - self.after_bytes

    @property
    def saved_display(self) -> str:
        """返回可读的节省空间描述，如 '减少 6.2MB'。"""
        saved = self.saved_bytes
        if saved >= 1024 * 1024:
            return f"减少 {saved / (1024 * 1024):.1f}MB"
        elif saved >= 1024:
            return f"减少 {saved / 1024:.0f}KB"
        return f"减少 {saved}B"

    def merge(self, other: ImageCompressResult) -> None:
        """合并另一个结果（用于多文档批处理）。"""
        self.processed += other.processed
        self.before_bytes += other.before_bytes
        self.after_bytes += other.after_bytes
        self.skipped += other.skipped
        self.errors.extend(other.errors)


# ============================================================
# 图片部件发现
# ============================================================


def _find_image_parts(doc: Document) -> list[Any]:
    """遍历文档包中的所有部件，返回 ImagePart 实例列表。

    通过 ``doc.part.package.parts`` 而非仅 ``doc.part.rels``，
    确保页眉、页脚等区域中的图片也被包含。
    """
    from docx.opc.part import Part  # lazy import to avoid circular ref

    image_parts: list[Any] = []
    # 收集所有匹配 partname 且是 ImagePart 类型的部件
    for part in doc.part.package.parts:
        # 通过 partname 后缀或 isinstance 判断
        if isinstance(part, Part) and hasattr(part, "blob"):
            # 检查是否图片部件（partname 以 /word/media/ 开头或 content_type 含 image）
            ct = getattr(part, "content_type", "") or ""
            pn = getattr(part, "partname", "") or ""
            if "image" in ct.lower() or "/media/" in str(pn):
                image_parts.append(part)
    return image_parts


# ============================================================
# blob 替换（封装私有属性）
# ============================================================


def replace_image_blob(part: Any, blob: bytes) -> None:
    """安全替换图片部件的二进制数据。

    封装对 ``_blob`` 私有属性的访问，隔离 python-docx 内部实现细节，
    未来升级时可在此处集中适配。
    """
    # python-docx 的 Blob 属性通过 _blob 字段存储
    # 目前没有公开的 setter，只能通过私有属性替换
    part._blob = blob
    # 清除 blob 缓存（如果存在），确保下次读取时返回新数据
    if hasattr(part, "_blob_cache"):
        del part._blob_cache


# ============================================================
# 单图压缩
# ============================================================


def _compress_single_image(blob: bytes, max_side: int, quality: int) -> bytes | None:
    """对单张图片执行缩放 + 压缩。

    规则:
      - 保持原始格式（不将 PNG 转 JPEG，保留透明通道）
      - 仅当最长边超过 ``max_side`` 时缩放
      - JPEG 使用 ``quality`` 参数; PNG 使用 optimize=True
      - 其他格式使用 ``optimize=True`` 减少体积

    返回压缩后的 bytes，失败返回 ``None``。
    """
    try:
        from PIL import Image as PILImage
    except ImportError:
        logger.warning("Pillow not installed, cannot compress images")
        return None

    try:
        with PILImage.open(io.BytesIO(blob)) as img:
            orig_fmt = img.format or "JPEG"
            w, h = img.size
            # 缩放
            if max(w, h) > max_side:
                ratio = max_side / max(w, h)
                new_w = int(w * ratio)
                new_h = int(h * ratio)
                img = img.resize((new_w, new_h), PILImage.LANCZOS)

            buf = io.BytesIO()
            save_kwargs: dict[str, Any] = {"optimize": True}

            if orig_fmt in ("JPEG", "JPG", "JFIF"):
                # JPEG: 指定 quality，确保 RGB 模式
                if img.mode != "RGB":
                    img = img.convert("RGB")
                save_kwargs["format"] = "JPEG"
                save_kwargs["quality"] = quality
            elif orig_fmt == "PNG":
                # PNG: 保留透明通道，不转 JPEG，只 optimize
                save_kwargs["format"] = "PNG"
            else:
                # GIF / WEBP / BMP 等 — 保持原格式
                save_kwargs["format"] = orig_fmt

            img.save(buf, **save_kwargs)
            return buf.getvalue()
    except Exception as exc:
        logger.warning("Failed to compress image: %s", exc)
        return None


# ============================================================
# 公开接口 — 压缩
# ============================================================


def compress_images(doc: Document, config: PictureConfig) -> ImageCompressResult:
    """对文档中所有嵌入图片执行压缩处理。

    依据 ``config.auto_compress`` 决定是否执行; 若未安装 Pillow
    则静默跳过并返回空结果。

    Args:
        doc: python-docx Document 实例。
        config: 图片格式化配置。

    Returns:
        处理统计结果（``ImageCompressResult``）。
    """
    result = ImageCompressResult()

    # 开关检查
    if not config.auto_compress:
        logger.info("Image compression disabled by config")
        return result

    # 检查 Pillow 是否可用
    try:
        from PIL import Image as PILImage  # noqa: F401 — 仅检测
    except ImportError:
        logger.warning(
            "Pillow is not installed. Install with: pip install Pillow>=10.0.0"
        )
        return result

    max_side = config.max_side_pixels
    quality = config.quality
    max_file_size = config.max_file_size

    image_parts = _find_image_parts(doc)
    if not image_parts:
        logger.info("No images found in document")
        return result

    logger.info(
        "Compressing %d image(s) (max_side=%dpx, quality=%d, max_file=%dB)",
        len(image_parts), max_side, quality, max_file_size,
    )

    for part in image_parts:
        blob = part.blob
        before_size = len(blob)

        # 检查是否需要压缩
        need_compress = False
        try:
            from PIL import Image as PILImage
            with PILImage.open(io.BytesIO(blob)) as img:
                w, h = img.size
                long_side = max(w, h)
                if long_side > max_side:
                    need_compress = True
                if before_size > max_file_size:
                    need_compress = True
        except Exception:
            result.errors.append(f"Failed to inspect image: {part.partname}")
            continue

        if not need_compress:
            result.skipped += 1
            continue

        # 执行压缩
        compressed = _compress_single_image(blob, max_side, quality)
        if compressed is not None:
            replace_image_blob(part, compressed)
            result.processed += 1
            result.before_bytes += before_size
            result.after_bytes += len(compressed)
            logger.debug(
                "Compressed %s: %dB -> %dB (%.1f%%)",
                part.partname, before_size, len(compressed),
                (1 - len(compressed) / before_size) * 100 if before_size else 0,
            )
        else:
            result.errors.append(f"Compression failed: {part.partname}")

    logger.info(
        "Image compression done: processed=%d, skipped=%d, %s",
        result.processed, result.skipped, result.saved_display,
    )
    return result


# ============================================================
# 图片尺寸设置
# ============================================================


def _unit_to_emu(value: float, unit: str) -> int:
    """将用户配置值转换为 EMU（English Metric Unit）。

    EMU 换算关系:
      1 cm = 360000 EMU
      1 mm = 36000 EMU
      1 pt = 12700 EMU
      1 inch = 914400 EMU
    """
    unit_map = {
        "cm": 360000, "厘米": 360000,
        "mm": 36000, "毫米": 36000,
        "pt": 12700, "磅": 12700,
        "inch": 914400, "英寸": 914400,
    }
    multiplier = unit_map.get(unit, 360000)  # 默认按 cm 换算
    return int(round(value * multiplier))


def _calculate_image_size(orig_cx: int, orig_cy: int,
                          config: PictureConfig) -> tuple[int | None, int | None]:
    """根据配置计算目标图片尺寸（EMU）。

    size_mode 控制策略:
      - "auto"   : 不修改，返回 (None, None)
      - "width"  : 按指定宽度，keep_ratio=True 时等比计算高度
      - "height" : 按指定高度，keep_ratio=True 时等比计算宽度

    no_enlarge 控制: 如果目标尺寸大于原始尺寸，则缩放到原始尺寸（而不是跳过）。

    Returns:
        (target_cx, target_cy) 或 (None, None) 如果无需修改。
    """
    size_mode = config.size_mode
    keep_ratio = config.keep_ratio
    no_enlarge = config.no_enlarge

    if size_mode == "auto":
        return None, None

    # 原始宽高比
    orig_ratio = orig_cx / orig_cy if orig_cy > 0 else 1.0

    if size_mode == "width":
        target_cx = _unit_to_emu(config.width, config.width_unit)
        if no_enlarge and target_cx > orig_cx:
            logger.debug("no_enlarge: clamp target %d to orig %d", target_cx, orig_cx)
            target_cx = orig_cx
        target_cy = int(round(target_cx / orig_ratio)) if keep_ratio else _unit_to_emu(config.height, config.height_unit)
        if no_enlarge and target_cy > orig_cy:
            target_cy = orig_cy
    elif size_mode == "height":
        target_cy = _unit_to_emu(config.height, config.height_unit)
        if no_enlarge and target_cy > orig_cy:
            logger.debug("no_enlarge: clamp target %d to orig %d", target_cy, orig_cy)
            target_cy = orig_cy
        target_cx = int(round(target_cy * orig_ratio)) if keep_ratio else _unit_to_emu(config.width, config.width_unit)
        if no_enlarge and target_cx > orig_cx:
            target_cx = orig_cx
    else:
        return None, None

    return target_cx, target_cy


def _resize_drawing_element(drawing, config: PictureConfig) -> None:
    """修改单个 ``w:drawing`` 元素中的图片尺寸。

    同时更新:
      - ``wp:extent``（容器级别，控制布局占位）
      - ``pic:spPr/a:xfrm/a:ext``（图片变换级别，控制渲染）
    """
    from docx.oxml.ns import qn

    # 查找容器（inline 或 anchor）
    container = drawing.find(qn("wp:inline"))
    if container is None:
        container = drawing.find(qn("wp:anchor"))
    if container is None:
        return

    # 获取当前尺寸
    extent = container.find(qn("wp:extent"))
    if extent is None:
        return

    try:
        orig_cx = int(extent.get("cx", "0"))
        orig_cy = int(extent.get("cy", "0"))
    except (ValueError, TypeError):
        return

    if orig_cx <= 0 or orig_cy <= 0:
        return

    # 计算目标尺寸
    target_cx, target_cy = _calculate_image_size(orig_cx, orig_cy, config)
    if target_cx is None or target_cx <= 0 or target_cy <= 0:
        return

    target_cx, target_cy = int(target_cx), int(target_cy)

    # 更新 wp:extent（容器级别）
    extent.set("cx", str(target_cx))
    extent.set("cy", str(target_cy))

    # 更新 pic:spPr/a:xfrm/a:ext（图片变换级别）
    for spPr in container.iter(qn("pic:spPr")):
        xfrm = spPr.find(qn("a:xfrm"))
        if xfrm is not None:
            a_ext = xfrm.find(qn("a:ext"))
            if a_ext is not None:
                a_ext.set("cx", str(target_cx))
                a_ext.set("cy", str(target_cy))

    logger.debug(
        "Image resized: %d x %d -> %d x %d EMU (no_enlarge=%s, size_mode=%s)",
        orig_cx, orig_cy, target_cx, target_cy,
        config.no_enlarge, config.size_mode,
    )


def apply_image_wrapping(doc: Document, config: PictureConfig) -> None:
    """设置文档中所有图片的文字环绕样式。

    根据 ``config.wrapping_style`` 将 ``<wp:inline>`` 转换为 ``<wp:anchor>``
    （非嵌入型环绕），或保持 ``<wp:inline>``（嵌入型）。

    Args:
        doc: python-docx Document 实例。
        config: 图片格式化配置（PictureConfig）。
    """
    from docx.oxml.ns import qn
    from lxml import etree

    wrapping_style = config.wrapping_style
    if wrapping_style == "inline":
        logger.debug("apply_image_wrapping: wrapping_style=inline, skip")
        return

    # 环绕类型 -> OOXML 标签映射
    wrap_tags = {
        "square": "wp:wrapSquare",
        "tight": "wp:wrapTight",
        "through": "wp:wrapThrough",
        "topBottom": "wp:wrapTopBottom",
        "behindText": "wp:wrapBehind",
        "inFrontOfText": "wp:wrapInFront",
    }

    wrap_tag = wrap_tags.get(wrapping_style)
    if wrap_tag is None:
        logger.warning("apply_image_wrapping: unknown wrapping_style=%s", wrapping_style)
        return

    nsmap = {
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }

    def _inline_to_anchor(inline_elem):
        """将单个 ``<wp:inline>`` 转换为 ``<wp:anchor>``。

        保留原有的 extent 和 graphic 子元素，根据环绕样式添加对应的 wrap 元素。

        注意：不同环绕类型需要不同的 wrap 元素结构：
          - square/tight/through: 有 wrapText 属性
          - topBottom: 无 wrapText，使用 distT/distB
          - behindText/inFrontOfText: 空元素，无属性
        """
        # 创建 wp:anchor 元素（直接放在 inline 之前）
        parent = inline_elem.getparent()
        anchor = etree.Element(qn("wp:anchor"))
        # 先插入 anchor 到 inline 之前，再删除 inline
        parent.insert(list(parent).index(inline_elem), anchor)
        parent.remove(inline_elem)

        # anchor 基本属性
        anchor.set("distT", inline_elem.get("distT", "0"))
        anchor.set("distB", inline_elem.get("distB", "0"))
        anchor.set("distL", inline_elem.get("distL", "0"))
        anchor.set("distR", inline_elem.get("distR", "0"))
        anchor.set("simplePos", "0")
        anchor.set("relativeHeight", "251658240")
        anchor.set("behindDoc", "1" if wrapping_style == "behindText" else "0")
        anchor.set("locked", "0")
        anchor.set("layoutInCell", "1")
        anchor.set("allowOverlap", "1")

        # 保留原有 extent
        extent = inline_elem.find(qn("wp:extent"))
        if extent is not None:
            anchor.append(extent)

        # effectExtent（零值）
        effect_extent = etree.SubElement(anchor, qn("wp:effectExtent"))
        effect_extent.set("l", "0")
        effect_extent.set("t", "0")
        effect_extent.set("r", "0")
        effect_extent.set("b", "0")

        # 根据环绕样式创建对应的 wrap 元素
        if wrapping_style == "behindText":
            # 空元素，无属性
            etree.SubElement(anchor, qn("wp:wrapBehind"))
        elif wrapping_style == "inFrontOfText":
            # 空元素，无属性
            etree.SubElement(anchor, qn("wp:wrapInFront"))
        elif wrapping_style == "topBottom":
            wrap_elem = etree.SubElement(anchor, qn("wp:wrapTopBottom"))
            wrap_elem.set("distT", "0")
            wrap_elem.set("distB", "0")
            wrap_elem.set("distL", "91440")  # ~0.25cm
            wrap_elem.set("distR", "91440")
        else:
            # square / tight / through — 有 wrapText 属性
            wrap_elem = etree.SubElement(anchor, qn(wrap_tag))
            wrap_elem.set("wrapText", "both")
            wrap_elem.set("distL", "91440")
            wrap_elem.set("distR", "91440")
            wrap_elem.set("distT", "0")
            wrap_elem.set("distB", "0")

        # 水平位置（相对于列）
        pos_h = etree.SubElement(anchor, qn("wp:positionH"))
        pos_h.set("relativeFrom", "column")
        pos_h_align = etree.SubElement(pos_h, qn("wp:align"))
        pos_h_align.text = "center"

        # 垂直位置（相对于段落）
        pos_v = etree.SubElement(anchor, qn("wp:positionV"))
        pos_v.set("relativeFrom", "paragraph")
        pos_v_offset = etree.SubElement(pos_v, qn("wp:posOffset"))
        pos_v_offset.text = "0"

        # 复制 graphic（包含图片数据引用）
        graphic = inline_elem.find(qn("a:graphic"))
        if graphic is not None:
            anchor.append(graphic)

        # 复制 docPr
        doc_pr = inline_elem.find(qn("wp:docPr"))
        if doc_pr is not None:
            anchor.append(doc_pr)

        # 复制 cNvGraphicFramePr
        cnv_graphic = inline_elem.find(qn("wp:cNvGraphicFramePr"))
        if cnv_graphic is not None:
            anchor.append(cnv_graphic)

        return anchor

    def _process_drawing(drawing):
        """处理单个 ``w:drawing`` 元素，转换 inline 为 anchor。"""
        inline = drawing.find(qn("wp:inline"))
        if inline is None:
            return  # 已经是 anchor 或其他类型，跳过
        _inline_to_anchor(inline)

    # 正文范围内
    body = doc.element.body
    for drawing in body.iter(qn("w:drawing")):
        _process_drawing(drawing)

    # 页眉/页脚
    for section in doc.sections:
        header = section.header
        if header and header.paragraphs:
            for para in header.paragraphs:
                for drawing in para._element.iter(qn("w:drawing")):
                    _process_drawing(drawing)
        footer = section.footer
        if footer and footer.paragraphs:
            for para in footer.paragraphs:
                for drawing in para._element.iter(qn("w:drawing")):
                    _process_drawing(drawing)

    logger.debug("apply_image_wrapping: wrapping_style=%s applied", wrapping_style)


def apply_image_alignment(doc: Document, config: PictureConfig) -> None:
    """设置文档中所有锚定图片的对齐方式。

    对于使用 ``<wp:anchor>`` 的图片，通过 ``positionH/align`` 控制水平对齐。
    对于使用 ``<wp:inline>`` 的图片，通过段落级 ``w:jc`` 控制对齐。

    Args:
        doc: python-docx Document 实例。
        config: 图片格式化配置（PictureConfig）。
    """
    from docx.oxml.ns import qn
    from lxml import etree

    alignment = config.alignment

    alignment_h_map = {
        "left": "left", "center": "center", "right": "right",
        "top": "center", "middle": "center", "bottom": "center",
        "distribute_h": "distribute", "distribute_v": "center",
    }
    target_h = alignment_h_map.get(alignment, "center")

    jc_map = {"left": "left", "center": "center", "right": "right"}

    def _set_anchor_alignment(anchor):
        """设置单个 ``<wp:anchor>`` 的对齐。"""
        pos_h = anchor.find(qn("wp:positionH"))
        if pos_h is None:
            pos_h = etree.SubElement(anchor, qn("wp:positionH"))
            pos_h.set("relativeFrom", "column")
        else:
            pos_h.set("relativeFrom", "column")
        for child in list(pos_h):
            pos_h.remove(child)
        h_align_elem = etree.SubElement(pos_h, qn("wp:align"))
        h_align_elem.text = target_h

        # 垂直对齐
        pos_v = anchor.find(qn("wp:positionV"))
        if pos_v is None:
            pos_v = etree.SubElement(anchor, qn("wp:positionV"))
            pos_v.set("relativeFrom", "paragraph")
        else:
            pos_v.set("relativeFrom", "paragraph")
        for child in list(pos_v):
            pos_v.remove(child)
        v_align_elem = etree.SubElement(pos_v, qn("wp:align"))
        v_align_elem.text = "center"

    def _set_inline_alignment(drawing):
        """通过段落对齐设置内联图片的对齐。"""
        if target_h not in jc_map:
            return
        # 找到包含此 drawing 的段落
        parent = drawing.getparent()
        while parent is not None:
            if parent.tag == qn("w:p"):
                # 设置段落对齐
                pPr = parent.find(qn("w:pPr"))
                if pPr is None:
                    pPr = etree.SubElement(parent, qn("w:pPr"))
                    parent.insert(0, pPr)
                # 移除旧的 jc
                old_jc = pPr.find(qn("w:jc"))
                if old_jc is not None:
                    pPr.remove(old_jc)
                jc = etree.SubElement(pPr, qn("w:jc"))
                jc.set(qn("w:val"), jc_map.get(target_h, "left"))
                return
            parent = parent.getparent()

    for drawing in doc.element.body.iter(qn("w:drawing")):
        anchor = drawing.find(qn("wp:anchor"))
        if anchor is not None:
            _set_anchor_alignment(anchor)
        else:
            _set_inline_alignment(drawing)

    # 页眉/页脚
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part and part.paragraphs:
                for para in part.paragraphs:
                    for drawing in para._element.iter(qn("w:drawing")):
                        anchor = drawing.find(qn("wp:anchor"))
                        if anchor is not None:
                            _set_anchor_alignment(anchor)
                        else:
                            _set_inline_alignment(drawing)

    logger.debug("apply_image_alignment: alignment=%s (h=%s) applied",
                 alignment, target_h)


def apply_image_size(doc: Document, config: PictureConfig) -> None:
    """遍历文档 XML 树中所有 ``w:drawing`` 元素，设置图片显示尺寸。

    通过 ``doc.element.body.iter()`` 全局搜索 ``w:drawing`` 元素，
    覆盖正文段落、表格单元格内图片。页眉/页脚图片仍通过 sections 单独遍历。

    根据 PictureConfig 的 size_mode / width / height / keep_ratio / no_enlarge
    调整文档中内联图片的渲染尺寸。

    Args:
        doc: python-docx Document 实例。
        config: 图片格式化配置（PictureConfig）。
    """
    if config.size_mode == "auto":
        logger.debug("apply_image_size: size_mode=auto, skip")
        return

    from docx.oxml.ns import qn

    # 正文范围内全局搜索 w:drawing（覆盖表格内图片）
    body = doc.element.body
    drawing_count = 0
    for drawing in body.iter(qn("w:drawing")):
        _resize_drawing_element(drawing, config)
        drawing_count += 1

    # 页眉/页脚中的图片
    for section in doc.sections:
        header = section.header
        if header and header.paragraphs:
            for para in header.paragraphs:
                for drawing in para._element.iter(qn("w:drawing")):
                    _resize_drawing_element(drawing, config)
                    drawing_count += 1
        footer = section.footer
        if footer and footer.paragraphs:
            for para in footer.paragraphs:
                for drawing in para._element.iter(qn("w:drawing")):
                    _resize_drawing_element(drawing, config)
                    drawing_count += 1

    logger.debug("apply_image_size: %d drawing(s) processed", drawing_count)

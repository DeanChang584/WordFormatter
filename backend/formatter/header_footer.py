"""
页眉页脚规则 — 字体、字号、字形、对齐方式

对文档所有 section 的页眉/页脚段落应用格式化。
"""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .font import apply_run_font

ALIGNMENT = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def apply_header_footer(doc, config) -> None:
    """将页眉/页脚字体配置应用到所有 section 的页眉和页脚段落。

    Args:
        doc: python-docx Document 实例。
        config: HeaderFooterConfig（shared.schemas）实例。
    """
    font_cn = config.font_cn
    font_en = config.font_en
    font_size = config.font_size
    font_style = config.font_style or "normal"
    font_bold = "bold" in font_style
    font_italic = "italic" in font_style
    alignment = ALIGNMENT.get(config.alignment, WD_ALIGN_PARAGRAPH.CENTER)

    for section in doc.sections:
        # Header
        header = section.header
        if header and not header.is_linked_to_previous:
            _format_paragraphs(header.paragraphs, font_cn, font_en,
                               font_size, font_bold, font_italic, alignment)

        # Footer
        footer = section.footer
        if footer and not footer.is_linked_to_previous:
            _format_paragraphs(footer.paragraphs, font_cn, font_en,
                               font_size, font_bold, font_italic, alignment)

        # Even-page header/footer (if present)
        even_header = section.even_page_header
        if even_header and not even_header.is_linked_to_previous:
            _format_paragraphs(even_header.paragraphs, font_cn, font_en,
                               font_size, font_bold, font_italic, alignment)
        even_footer = section.even_page_footer
        if even_footer and not even_footer.is_linked_to_previous:
            _format_paragraphs(even_footer.paragraphs, font_cn, font_en,
                               font_size, font_bold, font_italic, alignment)


def _format_paragraphs(paragraphs, font_cn, font_en, font_size,
                       font_bold, font_italic, alignment) -> None:
    for para in paragraphs:
        para.paragraph_format.alignment = alignment
        for run in para.runs:
            apply_run_font(run, font_cn, font_en, font_size,
                           "#000000", font_bold, font_italic)
        # Also set the paragraph-level default font via pPr/rPr
        _ensure_run_properties(para, font_cn, font_en, font_size,
                               font_bold, font_italic)


def _ensure_run_properties(para, font_cn, font_en, font_size,
                           font_bold, font_italic) -> None:
    """If the paragraph has no runs (e.g. empty header), still set
    a default run so that new content typed by the user inherits
    the configured font."""
    if para.runs:
        return
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import Pt

    pPr = para._element.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        pPr.append(rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_cn)
    rFonts.set(qn("w:ascii"), font_en)
    rFonts.set(qn("w:hAnsi"), font_en)

    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = parse_xml(f'<w:sz {nsdecls("w")}/>')
        rPr.append(sz)
    sz.set(qn("w:val"), str(int(font_size * 2)))  # half-points

    if font_bold:
        b = rPr.find(qn("w:b"))
        if b is None:
            b = parse_xml(f'<w:b {nsdecls("w")}/>')
            rPr.append(b)

    if font_italic:
        i = rPr.find(qn("w:i"))
        if i is None:
            i = parse_xml(f'<w:i {nsdecls("w")}/>')
            rPr.append(i)

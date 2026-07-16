"""
页面规则 — 纸张大小、方向、边距、页眉页脚边距、页码、自定义纸张、文档网格

Step 2.3 完整实现。
接受 FormatProfile.page（data_model.PageConfig）。
所有边距单位已标准化为 mm（转换在 engine.py 中完成）。
"""

from docx.shared import Mm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from lxml import etree
from .data_model import DocumentGridConfig, HeaderFooterConfig, PageConfig, PAPER_SIZES

# 最小合法边距（mm），防止 0/negative 导致 Word 报错
_MIN_MARGIN_MM = 0.1


def _resolve_paper_size(config: PageConfig) -> tuple[float, float]:
    """返回 (宽mm, 高mm)。

    - 标准尺寸从 PAPER_SIZES 字典查找（未知名称 fallback A4）
    - 自定义尺寸使用 custom_width / custom_height，< 1mm 时 fallback A4
    """
    if config.paper_size == "custom":
        w = config.custom_width
        h = config.custom_height
        if w >= 1.0 and h >= 1.0:
            return w, h
        # 无效自定义尺寸 → fallback A4
        return PAPER_SIZES["A4"]
    return PAPER_SIZES.get(config.paper_size, PAPER_SIZES["A4"])


def _add_page_number_footer(section) -> None:
    """向 section 的默认 footer 插入居中页码（PAGE field）。

    仅在 footer 尚无内容时插入，避免覆盖用户已有的页眉页脚。
    若 footer 已有段落则跳过。
    """
    footer = section.footer
    if footer.paragraphs and any(p.text.strip() for p in footer.paragraphs):
        return  # footer 已有内容，不覆盖

    # 确保 footer 不再链接到 previous section
    footer.is_linked_to_previous = False

    # 清除空段落，用一个新段落承载页码
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    para = footer.add_paragraph()
    para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

    # 插入 PAGE field code：fieldSimple 方式
    # <w:fldSimple w:instr=" PAGE ">
    #   <w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:eastAsia="宋体"
    #          w:hAnsi="Times New Roman"/><w:sz w:val="18"/></w:rPr>
    #     <w:t>1</w:t>
    #   </w:r>
    # </w:fldSimple>
    rpr_xml = (
        f'<w:rPr {nsdecls("w")}>'
        f'<w:rFonts w:ascii="Times New Roman" w:eastAsia="\u5b8b\u4f53"'
        f' w:hAnsi="Times New Roman"/>'
        f'<w:sz w:val="18"/>'   # 9pt = 18 half-points
        f'</w:rPr>'
    )
    fld_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
        f'  <w:r>{rpr_xml}<w:t>1</w:t></w:r>'
        f'</w:fldSimple>'
    )
    para._element.append(parse_xml(fld_xml))


def apply_document_grid(section, config: DocumentGridConfig,
                       page_width_mm: float = 210.0, page_height_mm: float = 297.0,
                       margin_top: float = 25.4, margin_bottom: float = 25.4,
                       margin_left: float = 31.7, margin_right: float = 31.7,
                       font_size_pt: float = 12.0) -> None:
    """将文档网格配置写入 section 的 XML（w:docGrid）。

    对应 Word 页面设置 → 文档网格。

    ``w:linePitch`` 用可用高度计算（Word 用 section 的可见行数反推、
    自动扣除边距），因此传入可用高度。``w:charSpace`` 同理用可用宽度，
    但受正文字号下界约束。

    Args:
        section: python-docx Section 对象
        config:  DocumentGridConfig 实例
        page_width_mm / page_height_mm: 纸张尺寸（mm）
        margin_*: 四边边距（mm）
        font_size_pt: 正文默认字号（pt），charSpace 下界
    """
    sectPr = section._sectPr

    # 移除已有的 w:docGrid（如有）
    for child in list(sectPr):
        if child.tag == qn('w:docGrid'):
            sectPr.remove(child)

    if config.mode == "none":
        return

    # 转换工具（1 mm ≈ 56.7 twips）
    MM_TO_TWIPS = 56.7

    available_w_twips = (page_width_mm - margin_left - margin_right) * MM_TO_TWIPS
    available_h_twips = (page_height_mm - margin_top - margin_bottom) * MM_TO_TWIPS

    # 正文字号下界（twips），Word 不允许单元格比字体更窄
    font_size_twips = font_size_pt * 20.0

    attrs = {}
    if config.mode == "lines":
        attrs['w:type'] = 'lines'
        attrs['w:linePitch'] = str(int(round(available_h_twips / config.lines_per_page)))
    elif config.mode == "both":
        # type="lines" avoids Word's linesAndChars font-size validation
        # that can silently override the computed values.
        attrs['w:type'] = 'lines'
        attrs['w:linePitch'] = str(int(round(available_h_twips / config.lines_per_page)))
        char_space = max(
            int(round(available_w_twips / config.chars_per_line)),
            int(font_size_twips),
        )
        attrs['w:charSpace'] = str(char_space)

    if config.adjust_right_indent:
        attrs['w:snapToGrid'] = '1'
    if config.align_to_grid:
        attrs['w:alignToGrid'] = '1'

    # lxml SubElement inherits the parent namespace context without
    # injecting extras — parse_xml and OxmlElement both add xmlns attrs
    # which Word may reject as conflicting with sectPr's own declarations.
    docGrid = etree.SubElement(sectPr, qn("w:docGrid"))
    for k, v in attrs.items():
        docGrid.set(qn(k), v)


def apply_page_setup(section, config: PageConfig, font_size_pt: float = 12.0) -> None:
    """将页面配置应用到文档 section。

    覆盖：边距、纸张大小（含自定义）、方向、页眉页脚边距、页码、文档网格。
    """
    # ── 边距（clamp 到最小值）──
    section.top_margin = Mm(max(_MIN_MARGIN_MM, config.margin_top))
    section.bottom_margin = Mm(max(_MIN_MARGIN_MM, config.margin_bottom))
    section.left_margin = Mm(max(_MIN_MARGIN_MM, config.margin_left))
    section.right_margin = Mm(max(_MIN_MARGIN_MM, config.margin_right))

    # ── 页眉页脚边距 ──
    section.header_distance = Mm(max(_MIN_MARGIN_MM, config.header_distance))
    section.footer_distance = Mm(max(_MIN_MARGIN_MM, config.footer_distance))

    # ── 纸张大小 + 方向 ──
    w_mm, h_mm = _resolve_paper_size(config)

    if config.orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        # 横向：宽 > 高
        section.page_width = Mm(max(w_mm, h_mm))
        section.page_height = Mm(min(w_mm, h_mm))
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        # 纵向：高 > 宽
        section.page_width = Mm(min(w_mm, h_mm))
        section.page_height = Mm(max(w_mm, h_mm))

    # ── 页码 ──
    if config.page_number:
        _add_page_number_footer(section)

    # ── 文档网格 ──
    if config.document_grid.mode != "none":
        apply_document_grid(section, config.document_grid,
                           page_width_mm=w_mm, page_height_mm=h_mm,
                           margin_top=config.margin_top,
                           margin_bottom=config.margin_bottom,
                           margin_left=config.margin_left,
                           margin_right=config.margin_right,
                           font_size_pt=font_size_pt)


def apply_header_footer_font(section, config: HeaderFooterConfig) -> None:
    """对 section 的页眉和页脚中的所有段落应用字体、字号、字形、对齐。

    遍历 header.paragraphs 和 footer.paragraphs，逐个 run 应用字体。
    若 footer 包含页码（fldSimple），同样更新其 run 字体。

    自动处理以下边缘情况：
    - header/footer 为空（没有段落）时创建默认段落
    - 段落中有 run 但没有字体设置时补全
    - 断开与上一节的链接（is_linked_to_previous = False）
    """
    # 解析字形
    style = config.font_style.lower() if config.font_style else "normal"
    bold = "bold" in style
    italic = "italic" in style

    # 对齐映射
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    alignment = align_map.get(config.alignment, WD_ALIGN_PARAGRAPH.CENTER)

    # 处理 header 和 footer
    for part in (section.header, section.footer):
        # 断开链接到上一节，确保设置独立生效
        part.is_linked_to_previous = False

        paragraphs = part.paragraphs

        # 如果没有任何段落，创建一个带空 run 的默认段落。
        # 页码通过 fldSimple 插入，没有 run 且 text 为空，
        # 需额外检查 fldSimple 以免误删原有页脚内容。
        def _has_content(p):
            if p.text.strip() or p.runs:
                return True
            if p._element.find(qn('w:fldSimple')) is not None:
                return True
            if p._element.find(qn('w:fldChar')) is not None:
                return True
            return False

        if not paragraphs or not any(_has_content(p) for p in paragraphs):
            # 清除已有的空段落元素
            for p in paragraphs:
                p._element.getparent().remove(p._element)
            part.add_paragraph().add_run("")

        # 再次获取段落列表（因为上面可能新增了段落）
        for para in part.paragraphs:
            # 清零缩进（避免继承原文档页眉页脚的缩进）
            _zero_paragraph_indent(para)

            # 设置段落对齐
            para.alignment = alignment

            # 遍历所有 run 应用字体
            for run in para.runs:
                run.font.size = Pt(config.font_size)
                run.font.bold = bold
                run.font.italic = italic
                run.font.name = config.font_en

                # 设置中文字体（XML 级）
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                    rPr.insert(0, rFonts)
                rFonts.set(qn("w:eastAsia"), config.font_cn)
                rFonts.set(qn("w:ascii"), config.font_en)
                rFonts.set(qn("w:hAnsi"), config.font_en)
                rFonts.set(qn("w:cs"), config.font_en)

            # 更新页码 fldSimple 中的 run（如果存在）
            for fld in para._element.findall(qn("w:fldSimple")):
                for r in fld.findall(qn("w:r")):
                    rPr = r.find(qn("w:rPr"))
                    if rPr is None:
                        rPr = parse_xml(f'<w:rPr {nsdecls("w")} />')
                        r.insert(0, rPr)

                    # 更新字号
                    sz = rPr.find(qn("w:sz"))
                    if sz is None:
                        sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="{int(config.font_size * 2)}"/>')
                        rPr.append(sz)
                    else:
                        sz.set(qn("w:val"), str(int(config.font_size * 2)))

                    # 更新字体
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is None:
                        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                        rPr.insert(0, rFonts)
                    rFonts.set(qn("w:eastAsia"), config.font_cn)
                    rFonts.set(qn("w:ascii"), config.font_en)
                    rFonts.set(qn("w:hAnsi"), config.font_en)
                    rFonts.set(qn("w:cs"), config.font_en)


def _zero_paragraph_indent(paragraph) -> None:
    """显式写入零值缩进，清除从原文档继承的缩进。"""
    pPr = paragraph._element.get_or_add_pPr()
    for old in pPr.findall(qn("w:ind")):
        pPr.remove(old)
    ind = parse_xml(
        f'<w:ind {nsdecls("w")} w:firstLine="0" w:firstLineChars="0" '
        f'w:hanging="0" w:hangingChars="0" w:left="0" w:right="0"/>')
    pPr.append(ind)

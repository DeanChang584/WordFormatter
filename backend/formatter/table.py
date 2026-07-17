"""
表格排版引擎 — 对 docx 文档中的所有表格应用格式化规则

支持的功能（匹配 shared.schemas.TableConfig）：
  - 表格对齐（整体居中/左/右）
  - 表格宽度（自动/固定）
  - 表头字体、字号、粗细、斜体、下划线、背景色
  - 边框样式（全部/无/仅横向/网格）、颜色、宽度
  - 单元格对齐（水平 + 垂直）
  - 单元格边距
  - 行高（自动/固定/最小）
  - 跨页断行 / 表头跨页重复
"""

from typing import Optional, TYPE_CHECKING

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from lxml import etree
from docx.shared import RGBColor, Pt

from .data_model import TableConfig as DmTableConfig

if TYPE_CHECKING:
    from docx.document import Document
else:
    Document = None  # runtime only used for type hint, not instantiated here


# ============================================================
# 辅助常量
# ============================================================

# python-docx WD_TABLE_ALIGNMENT 映射
_ALIGN_MAP = {
    "left": WD_TABLE_ALIGNMENT.LEFT,
    "center": WD_TABLE_ALIGNMENT.CENTER,
    "right": WD_TABLE_ALIGNMENT.RIGHT,
}

# 单元格水平对齐（段落级别）
_CELL_ALIGN_H = {
    "left": "left",
    "center": "center",
    "right": "right",
}

# 单元格垂直对齐（表格属性级别）
_CELL_ALIGN_V = {
    "top": "top",
    "middle": "center",
    "bottom": "bottom",
}


def _emu_from_cm(value: float) -> int:
    """厘米 → EMU（OpenXML 单位）.  1 cm = 360 000 EMU."""
    return int(value * 360000)


def _emu_from_pt(value: float) -> int:
    """磅 → EMU"""
    return int(value * 12700)


def _twips_from_indent(value: float, unit: str) -> int:
    """将缩进值转换为 twips（1/20 磅）。

    支持的输入单位：
        "字符" — 1 字符 ≈ 240 twips (约 4.233 mm)
        "厘米" — 1 cm = 567 twips
        "毫米" — 1 mm = 56.7 twips
    """
    if unit == "字符":
        return int(value * 240)
    if unit == "毫米":
        return int(value * 56.7)
    # 厘米（默认）
    return int(value * 567)


def _rgb_from_hex(hex_color: str) -> Optional[RGBColor]:
    """#RRGGBB → RGBColor 对象, 或 None（当为 'none' 时）"""
    if not hex_color or hex_color == "none":
        return None
    h = hex_color.lstrip("#")
    if len(h) != 6:
        return None
    return RGBColor(
        int(h[0:2], 16),
        int(h[2:4], 16),
        int(h[4:6], 16),
    )


# ============================================================
# 核心格式化函数
# ============================================================


def apply_table_format(doc: "Document", config: DmTableConfig) -> None:
    """对文档中所有表格（包括嵌套表格）应用格式化规则。"""
    for table in doc.tables:
        _apply_table_recursive(table, config)


def _apply_table_recursive(table, config: DmTableConfig) -> None:
    """对单个表格及其所有嵌套子表格递归应用格式化规则。"""
    _set_table_alignment(table, config)
    _set_table_width(table, config)
    _set_borders(table, config)
    _set_cell_margins(table, config)

    for row_idx, row in enumerate(table.rows):
        # 行高
        _set_row_height(row, config, row_idx)

        # 跨页断行（允许行跨页）
        _set_row_allow_split(row, config)

        for cell in row.cells:
            # 表头格式
            if row_idx == 0:
                _set_header_format(cell, config)
            else:
                _set_body_format(cell, config)

            # 单元格对齐
            _set_cell_alignment(cell, config)

            # 行距
            _apply_cell_line_spacing(cell, config)

            # 递归处理嵌套表格
            for nested_table in _get_nested_tables(cell):
                _apply_table_recursive(nested_table, config)

        # 表头文字居中覆盖：头行水平对齐强制为居中
        if row_idx == 0 and config.header_text_center:
            for cell in row.cells:
                _override_header_cell_align_center(cell)

    # 表头跨页重复
    if config.repeat_header:
        _set_repeat_header(table)


def _get_nested_tables(cell):
    """返回单元格内所有直接嵌套的 Table 对象。"""
    from docx.table import Table
    nested = []
    tc = cell._tc  # noqa: SLF001
    for tbl_elem in tc.findall(qn("w:tbl")):
        nested.append(Table(tbl_elem, None))
    return nested


# ============================================================
# 内部实现
# ============================================================


def _set_table_alignment(table, config: DmTableConfig):
    """整体表格对齐。"""
    alignment = _ALIGN_MAP.get(config.table_alignment)
    if alignment is not None:
        table.alignment = alignment


def _set_table_width(table, config: DmTableConfig):
    """表格宽度（固定值或自动）。"""
    tbl = table._tbl  # noqa: SLF001
    tbl_pr = tbl.tblPr

    if config.width_mode == "fixed":
        # 删除可能存在的 auto fit 属性
        for auto_fit in tbl_pr.findall(qn("w:tblW")):
            tbl_pr.remove(auto_fit)

        if config.width_unit == "%":
            # 百分比：OpenXML 中 5000 = 100%，1% = 50
            pct_val = int(config.width_value * 50)
            tbl_width = parse_xml(
                f'<w:tblW {nsdecls("w")} w:w="{pct_val}" w:type="pct"/>'
            )
        elif config.width_unit == "mm":
            # 毫米 → EMU (1 mm = 36000 EMU)
            width_emu = int(config.width_value * 36000)
            tbl_width = parse_xml(
                f'<w:tblW {nsdecls("w")} w:w="{width_emu}" w:type="dxa"/>'
            )
        else:
            # 厘米（默认）
            width_emu = _emu_from_cm(config.width_value)
            tbl_width = parse_xml(
                f'<w:tblW {nsdecls("w")} w:w="{width_emu}" w:type="dxa"/>'
            )
        tbl_pr.append(tbl_width)

        # 设置列宽为固定模式（禁止自动适应）
        _set_autofit(table, False)
    else:
        # 自动模式：设置自动适应
        _set_autofit(table, True)

        # 设置自动宽度（100%）
        for auto_fit in tbl_pr.findall(qn("w:tblW")):
            tbl_pr.remove(auto_fit)
        tbl_width = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
        tbl_pr.append(tbl_width)


def _set_autofit(table, enabled: bool):
    """设置表格自动适应内容。"""
    tbl = table._tbl  # noqa: SLF001
    tbl_pr = tbl.tblPr

    # 删除旧的 autofit 属性
    for old in tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(old)

    if not enabled:
        # 固定列宽
        layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    else:
        # 自动适应
        layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="autofit"/>')
    tbl_pr.append(layout)


def _set_row_height(row, config: DmTableConfig, row_idx: int):
    """设置行高。"""
    tr = row._tr  # noqa: SLF001
    tr_pr = tr.find(qn("w:trPr"))

    mode = config.row_height_mode
    height_val = config.row_height

    # EMU 转换
    if config.row_height_unit == "cm" or config.row_height_unit == "厘米":
        height_emu = _emu_from_cm(height_val)
    elif config.row_height_unit == "mm" or config.row_height_unit == "毫米":
        height_emu = int(height_val * 36000)  # 1 mm = 36000 EMU
    elif config.row_height_unit == "pt" or config.row_height_unit == "磅":
        height_emu = _emu_from_pt(height_val)
    else:
        height_emu = _emu_from_cm(height_val)  # 默认按 cm

    # 构建行高 XML
    if mode == "fixed":
        # 固定高度（exact）
        rule = parse_xml(
            f'<w:trHeight {nsdecls("w")} w:val="{height_emu}" w:hRule="exact"/>'
        )
    elif mode == "at_least":
        # 最小高度
        rule = parse_xml(
            f'<w:trHeight {nsdecls("w")} w:val="{height_emu}" w:hRule="atLeast"/>'
        )
    else:
        # 自动 — 删除所有行高约束，让内容自然撑开
        if tr_pr is not None:
            for old in tr_pr.findall(qn("w:trHeight")):
                tr_pr.remove(old)
        return

    # 删除旧的行高
    if tr_pr is not None:
        for old in tr_pr.findall(qn("w:trHeight")):
            tr_pr.remove(old)
        tr_pr.append(rule)
    else:
        # 如果没有 trPr，创建它
        tr_pr_xml = f'<w:trPr {nsdecls("w")}/>'
        new_tr_pr = parse_xml(tr_pr_xml)
        new_tr_pr.append(rule)
        tr.insert(0, new_tr_pr)


def _set_row_allow_split(row, config: DmTableConfig):
    """设置行是否可以跨页断行。"""
    tr = row._tr  # noqa: SLF001
    tr_pr = tr.find(qn("w:trPr"))

    if config.auto_split:
        # 允许跨页断行：删除 can't split 属性
        if tr_pr is not None:
            for old in tr_pr.findall(qn("w:cantSplit")):
                tr_pr.remove(old)
    else:
        # 禁止跨页断行
        if tr_pr is None:
            tr_pr_xml = f'<w:trPr {nsdecls("w")}/>'
            tr_pr = parse_xml(tr_pr_xml)
            tr.insert(0, tr_pr)
        elif tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def _set_borders(table, config: DmTableConfig):
    """设置表格边框 — 使用单元格级边框（w:tcBorders）确保兼容性。

    在 OOXML 中，单元格级边框（w:tcBorders）优先级高于表格级边框（w:tblBorders），
    且某些 Word/WPS 版本对表格级边框的渲染存在兼容性问题。
    因此采用单元格级边框方案，为每个单元格的四个边设置显式边框。

    同时删除表格样式引用 (w:tblStyle)，防止样式定义的边框干扰。
    """
    tbl = table._tbl  # noqa: SLF001
    tbl_pr = tbl.tblPr

    # 删除旧表格级边框
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)

    # 删除所有表格属性级格式覆盖（样式引用、外观设定、主题等）
    for old in tbl_pr.findall(qn("w:tblStyle")):
        tbl_pr.remove(old)
    for old in tbl_pr.findall(qn("w:tblLook")):
        tbl_pr.remove(old)
    for old in tbl_pr.findall(qn("w:tblStyleRowBandSize")):
        tbl_pr.remove(old)
    for old in tbl_pr.findall(qn("w:tblStyleColBandSize")):
        tbl_pr.remove(old)

    style = config.border_style
    if style == "none":
        # 无边框：清除所有单元格级边框
        _clear_cell_borders(table)
        return

    # 构建边框属性
    color = config.border_color.replace("#", "")
    # w:sz 以 1/8 磅为单位。例如 0.5pt → 4，1pt → 8，2pt → 16
    width_eighths = int(round(config.border_width * 8))
    width_val = str(max(width_eighths, 2))  # 最小 2 (0.25pt)，过细可能不可见

    # 为每个单元格设置边框
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, style, color, width_val)


def _set_cell_borders(cell, style: str, color: str, width_val: str):
    """为单个单元格设置边框（w:tcBorders）。"""
    tc = cell._tc  # noqa: SLF001
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
        tc.insert(0, tc_pr)

    # 删除旧单元格级边框
    for old in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(old)

    if style == "all":
        # 全部框线 — 上下左右四条边
        borders_xml = f"""<w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="{width_val}" w:color="{color}" w:space="0"/>
            <w:left w:val="single" w:sz="{width_val}" w:color="{color}" w:space="0"/>
            <w:bottom w:val="single" w:sz="{width_val}" w:color="{color}" w:space="0"/>
            <w:right w:val="single" w:sz="{width_val}" w:color="{color}" w:space="0"/>
        </w:tcBorders>"""
    elif style == "horizontal":
        # 仅横向框线 — 上下两条边，左右显式设为无（防止默认样式继承导致出现竖向框线）
        borders_xml = f"""<w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="{width_val}" w:color="{color}" w:space="0"/>
            <w:left w:val="none" w:space="0"/>
            <w:bottom w:val="single" w:sz="{width_val}" w:color="{color}" w:space="0"/>
            <w:right w:val="none" w:space="0"/>
        </w:tcBorders>"""
    elif style == "grid":
        # 网格线 — 上下左右四条边
        borders_xml = f"""<w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="{width_val}" w:color="{color}" w:space="0"/>
            <w:left w:val="single" w:sz="{width_val}" w:color="{color}" w:space="0"/>
            <w:bottom w:val="single" w:sz="{width_val}" w:color="{color}" w:space="0"/>
            <w:right w:val="single" w:sz="{width_val}" w:color="{color}" w:space="0"/>
        </w:tcBorders>"""
    else:
        return

    borders_el = parse_xml(borders_xml)
    tc_pr.append(borders_el)


def _clear_cell_borders(table):
    """清除所有单元格级别的边框（w:tcBorders），但保留底纹。
    
    ！！注意！！此函数只清除边框不清除底纹，确保背景色设置不受影响。
    """
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc  # noqa: SLF001
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is not None:
                for old in tc_pr.findall(qn("w:tcBorders")):
                    tc_pr.remove(old)


def _insert_in_order(parent, el, successor_tags):
    """按 OOXML schema 顺序插入子元素：插在第一个存在的后继元素之前，否则追加到末尾。"""
    for tag in successor_tags:
        anchor = parent.find(qn(tag))
        if anchor is not None:
            anchor.addprevious(el)
            return
    parent.append(el)


def _set_cell_margins(table, config: DmTableConfig):
    """设置单元格边距。

    水平边距（left/right）由 cell_margin_h 控制，
    垂直边距（top/bottom）由 cell_margin_v 控制。

    同时写入两个级别：
      1. 表格级 w:tblCellMar（按 schema 顺序插入 tblPr，作为默认值）
      2. 单元格级 w:tcPr/w:tcMar（优先级最高，覆盖表格样式继承的边距）
    """
    # Convert horizontal margin to twips
    if config.cell_margin_h_unit == "mm" or config.cell_margin_h_unit == "毫米":
        margin_h_val = str(int(config.cell_margin_h * 56.7))  # mm → twips
    else:
        margin_h_val = str(int(config.cell_margin_h * 567))   # cm → twips

    # Convert vertical margin to twips
    if config.cell_margin_v_unit == "mm" or config.cell_margin_v_unit == "毫米":
        margin_v_val = str(int(config.cell_margin_v * 56.7))  # mm → twips
    else:
        margin_v_val = str(int(config.cell_margin_v * 567))   # cm → twips

    def _margins_xml(root_tag: str) -> str:
        return f"""<w:{root_tag} {nsdecls("w")}>
        <w:top w:w="{margin_v_val}" w:type="dxa"/>
        <w:left w:w="{margin_h_val}" w:type="dxa"/>
        <w:bottom w:w="{margin_v_val}" w:type="dxa"/>
        <w:right w:w="{margin_h_val}" w:type="dxa"/>
    </w:{root_tag}>"""

    # 1. 表格级默认边距 — 必须插在 w:tblLook 之前（schema 顺序），否则 Word 忽略
    tbl = table._tbl  # noqa: SLF001
    tbl_pr = tbl.tblPr

    for old in tbl_pr.findall(qn("w:tblCellMar")):
        tbl_pr.remove(old)

    margins_el = parse_xml(_margins_xml("tblCellMar"))
    _insert_in_order(tbl_pr, margins_el,
                     ("w:tblLook", "w:tblCaption", "w:tblDescription"))

    # 2. 单元格级边距 — 覆盖原文档单元格/表格样式自带的 tcMar
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc  # noqa: SLF001
            tc_pr = tc.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                tc.insert(0, tc_pr)

            for old in tc_pr.findall(qn("w:tcMar")):
                tc_pr.remove(old)

            tc_mar = parse_xml(_margins_xml("tcMar"))
            # schema 顺序：tcMar 必须在 textDirection/tcFitText/vAlign/hideMark 之前
            _insert_in_order(tc_pr, tc_mar,
                             ("w:textDirection", "w:tcFitText",
                              "w:vAlign", "w:hideMark"))


def _apply_font_rpr(run, config: DmTableConfig):
    """对单个 run 应用全局字形（加粗 / 斜体 / 下划线）。"""
    rpr = run._r.find(qn("w:rPr"))  # noqa: SLF001
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}/>')  # noqa: SLF001
        run._r.insert(0, rpr)  # noqa: SLF001

    # 加粗 — 显式设置 w:val，避免从样式继承
    b_el = rpr.find(qn("w:b"))
    if config.font_bold:
        if b_el is None:
            rpr.append(parse_xml(f'<w:b {nsdecls("w")} w:val="true"/>'))
        else:
            b_el.set(qn("w:val"), "true")
    else:
        if b_el is None:
            rpr.append(parse_xml(f'<w:b {nsdecls("w")} w:val="false"/>'))
        else:
            b_el.set(qn("w:val"), "false")

    # 斜体
    i_el = rpr.find(qn("w:i"))
    if config.font_italic:
        if i_el is None:
            rpr.append(parse_xml(f'<w:i {nsdecls("w")}/>'))
    else:
        if i_el is not None:
            rpr.remove(i_el)

    # 下划线
    u_el = rpr.find(qn("w:u"))
    if config.font_underline:
        if u_el is None:
            rpr.append(
                parse_xml(f'<w:u {nsdecls("w")} w:val="single"/>')
            )
    else:
        if u_el is not None:
            rpr.remove(u_el)


def _apply_header_font_rpr(run, config: DmTableConfig):
    """对单个 run 应用表头专属字形（仅加粗，斜体/下划线由全局字形控制）。"""
    rpr = run._r.find(qn("w:rPr"))  # noqa: SLF001
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}/>')  # noqa: SLF001
        run._r.insert(0, rpr)  # noqa: SLF001

    # 加粗（表头单独控制，覆盖全局字形设置）
    b_el = rpr.find(qn("w:b"))
    if config.header_bold:
        if b_el is None:
            rpr.append(parse_xml(f'<w:b {nsdecls("w")} w:val="true"/>'))
        else:
            b_el.set(qn("w:val"), "true")
    else:
        if b_el is None:
            rpr.append(parse_xml(f'<w:b {nsdecls("w")} w:val="false"/>'))
        else:
            b_el.set(qn("w:val"), "false")


def _set_header_format(cell, config: DmTableConfig):
    """设置表头（第一行）单元格格式：全局字形 + 表头专属字形 + 表头字体/字号。"""
    for para in cell.paragraphs:
        for run in para.runs:
            # 1. 全局字形
            _apply_font_rpr(run, config)

            # 2. 表头专属字形（叠加，覆盖全局）
            _apply_header_font_rpr(run, config)

            # 3. 表头字体 / 字号
            rpr = run._r.find(qn("w:rPr"))  # noqa: SLF001
            if rpr is None:
                rpr = parse_xml(f'<w:rPr {nsdecls("w")}/>')  # noqa: SLF001
                run._r.insert(0, rpr)  # noqa: SLF001

            # 中文字体
            font_cn = rpr.find(qn("w:rFonts"))
            if font_cn is None:
                font_cn = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                rpr.append(font_cn)
            font_cn.set(qn("w:eastAsia"), config.header_font_cn)
            font_cn.set(qn("w:ascii"), config.header_font_en)
            font_cn.set(qn("w:hAnsi"), config.header_font_en)

            # 字号（pt → half-pt）
            sz = rpr.find(qn("w:sz"))
            if sz is None:
                sz = parse_xml(f'<w:sz {nsdecls("w")}/>')
                rpr.append(sz)
            sz.set(qn("w:val"), str(int(config.header_size * 2)))

    # 表头背景色
    # 兼容 ""（前端默认）和 "none"（Python 默认）均表示不设置背景色
    if config.header_bg_color and config.header_bg_color != "none":
        tc = cell._tc  # noqa: SLF001
        tc_pr = tc.find(qn("w:tcPr"))
        if tc_pr is None:
            tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
            tc.insert(0, tc_pr)

        # 删除旧底纹
        for old in tc_pr.findall(qn("w:shd")):
            tc_pr.remove(old)

        color = config.header_bg_color.lstrip("#")
        shd = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
        )
        # schema 顺序：shd 必须在 noWrap/tcMar/vAlign 之前
        _insert_in_order(tc_pr, shd,
                         ("w:noWrap", "w:tcMar", "w:textDirection",
                          "w:tcFitText", "w:vAlign", "w:hideMark"))


def _set_body_format(cell, config: DmTableConfig):
    """设置正文单元格格式：应用全局字形 + 字体/字号（与表头使用相同设置）+ 背景色。"""
    for para in cell.paragraphs:
        for run in para.runs:
            # 1. 全局字形
            _apply_font_rpr(run, config)

            # 2. 字体族 + 字号（与表头相同的字体设置）
            rpr = run._r.find(qn("w:rPr"))
            if rpr is None:
                rpr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
                run._r.insert(0, rpr)

            # 中文字体 / 西文字体
            font = rpr.find(qn("w:rFonts"))
            if font is None:
                font = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                rpr.append(font)
            font.set(qn("w:eastAsia"), config.header_font_cn)
            font.set(qn("w:ascii"), config.header_font_en)
            font.set(qn("w:hAnsi"), config.header_font_en)

            # 字号（pt → half-pt）
            sz = rpr.find(qn("w:sz"))
            if sz is None:
                sz = parse_xml(f'<w:sz {nsdecls("w")}/>')
                rpr.append(sz)
            sz.set(qn("w:val"), str(int(config.header_size * 2)))

    # 正文背景色（与表头相同的逻辑）
    if config.body_bg_color and config.body_bg_color != "none":
        tc = cell._tc  # noqa: SLF001
        tc_pr = tc.find(qn("w:tcPr"))
        if tc_pr is None:
            tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
            tc.insert(0, tc_pr)

        # 删除旧底纹
        for old in tc_pr.findall(qn("w:shd")):
            tc_pr.remove(old)

        color = config.body_bg_color.lstrip("#")
        shd = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
        )
        # schema 顺序：shd 必须在 noWrap/tcMar/vAlign 之前
        _insert_in_order(tc_pr, shd,
                         ("w:noWrap", "w:tcMar", "w:textDirection",
                          "w:tcFitText", "w:vAlign", "w:hideMark"))


def _set_cell_alignment(cell, config: DmTableConfig):
    """设置单元格对齐。"""
    # 垂直对齐
    align_v = _CELL_ALIGN_V.get(config.cell_align_v, "center")

    # 获取或创建 tcPr
    tc = cell._tc  # noqa: SLF001
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
        tc.insert(0, tc_pr)

    # 垂直对齐
    for old in tc_pr.findall(qn("w:vAlign")):
        tc_pr.remove(old)
    valign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align_v}"/>')
    tc_pr.append(valign)

    # 水平对齐（段落级别）
    align_h = _CELL_ALIGN_H.get(config.cell_align_h, "center")
    for para in cell.paragraphs:
        ppr = para._p.find(qn("w:pPr"))  # noqa: SLF001
        if ppr is None:
            ppr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            para._p.insert(0, ppr)  # noqa: SLF001

        # 特殊格式（缩进）— 必须在 jc 之前（OOXML schema 顺序）
        for old in ppr.findall(qn("w:ind")):
            ppr.remove(old)
        ind_el = OxmlElement("w:ind")

        indent_type = config.indent_type
        if indent_type == "first_line":
            ind_el.set(qn("w:firstLine"),
                       str(_twips_from_indent(config.indent_value, config.indent_unit)))
        elif indent_type == "hanging":
            ind_el.set(qn("w:hanging"),
                       str(_twips_from_indent(config.indent_value, config.indent_unit)))
        elif indent_type != "none":
            ind_el.set(qn("w:left"),
                       str(_twips_from_indent(config.indent_value, config.indent_unit)))
        else:
            ind_el.set(qn("w:firstLine"), "0")
            ind_el.set(qn("w:firstLineChars"), "0")
            ind_el.set(qn("w:hanging"), "0")
            ind_el.set(qn("w:hangingChars"), "0")
            ind_el.set(qn("w:left"), "0")
            ind_el.set(qn("w:right"), "0")
        ppr.append(ind_el)

        # 对齐方式
        for old in ppr.findall(qn("w:jc")):
            ppr.remove(old)
        jc = parse_xml(f'<w:jc {nsdecls("w")} w:val="{align_h}"/>')
        ppr.append(jc)


def _pt_from_unit(value: float, unit: str) -> float:
    """将值转换为磅（pt）。

    Support:
        "pt" — 原值
        "cm" — 1 cm ≈ 28.3465 pt
        "mm" — 1 mm ≈ 2.83465 pt
    """
    if unit == "cm":
        return value * 28.3465
    if unit == "mm":
        return value * 2.83465
    return value  # pt (default)


def _apply_cell_line_spacing(cell, config: DmTableConfig):
    """Apply line spacing to all paragraphs in a table cell using raw XML.

    Directly writes w:spacing/w:line + w:lineRule onto the paragraph element,
    bypassing python-docx paragraph_format.line_spacing which is unreliable
    for paragraphs inside table cells.

    OpenXML logic:
        multiple (auto): w:line = value × 240 twips,  w:lineRule = "auto"
        fixed (exact):   w:line = pt × 20 twips,       w:lineRule = "exact"
        at_least:        w:line = pt × 20 twips,       w:lineRule = "atLeast"
    """
    mode = config.line_spacing_mode
    value = config.line_spacing
    unit = getattr(config, "line_spacing_unit", "pt")

    if mode == "fixed":
        line_twips = str(int(round(_pt_from_unit(value, unit) * 20)))
        line_rule = "exact"
    elif mode == "at_least":
        line_twips = str(int(round(_pt_from_unit(value, unit) * 20)))
        line_rule = "atLeast"
    else:  # multiple
        line_twips = str(int(round(value * 240)))
        line_rule = "auto"

    for para in cell.paragraphs:
        pPr = para._element.get_or_add_pPr()
        # 清除旧的 w:spacing 元素
        old_sp = pPr.find(qn("w:spacing"))
        if old_sp is not None:
            pPr.remove(old_sp)
        # 新建 w:spacing 并用 line/lineRule 替换
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:line"), line_twips)
        sp.set(qn("w:lineRule"), line_rule)
        # 将 w:spacing 插入到 w:ind 之前（OOXML 规范要求 w:spacing 在 w:ind 之前）
        # 否则 Word 可能忽略行距设置
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            pPr.insert(pPr.index(ind), sp)
        else:
            pPr.append(sp)


def _override_header_cell_align_center(cell):
    """将表头单元格的水平对齐强制覆盖为居中（用于 header_text_center 选项）。"""
    for para in cell.paragraphs:
        ppr = para._p.find(qn("w:pPr"))
        if ppr is None:
            ppr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
            para._p.insert(0, ppr)

        # 删除旧对齐
        for old in ppr.findall(qn("w:jc")):
            ppr.remove(old)
        # 写入居中
        jc = parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>')
        ppr.append(jc)


def _set_repeat_header(table):
    """设置表头行在跨页时重复。"""
    for row in table.rows:
        if row == table.rows[0]:
            tr = row._tr  # noqa: SLF001
            tr_pr = tr.find(qn("w:trPr"))
            if tr_pr is None:
                tr_pr = parse_xml(f'<w:trPr {nsdecls("w")}/>')
                tr.insert(0, tr_pr)

            # 删除旧的 header 标记
            for old in tr_pr.findall(qn("w:tblHeader")):
                tr_pr.remove(old)

            header = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
            tr_pr.append(header)
            break  # 只需要处理第一行


def _is_in_table_cell(paragraph) -> bool:
    """检查段落是否位于表格单元格内（用于引擎跳过表格段落）。

    通过检查段落 XML 元素的父级链中是否包含 w:tc 标签来判断。
    如果段落位于表格单元格内，则该段落应由 apply_table_format 统一控制，
    不应在引擎的逐段落覆盖步骤中处理。
    """
    parent = paragraph._element.getparent()
    while parent is not None:
        tag = parent.tag.split("}")[-1] if "}" in parent.tag else parent.tag
        if tag == "tc":
            return True
        parent = parent.getparent()
    return False
